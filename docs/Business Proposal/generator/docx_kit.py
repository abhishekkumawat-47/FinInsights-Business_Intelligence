"""The typographic system for the proposal, expressed once as reusable builders.

Every measurement here is a decision, not a default:

* Body 10.5 pt on 130% leading, in a 15 cm measure, inside Butterick's 10 to 12 pt,
  120 to 145% and 45 to 90 character bands (Practical Typography, key rules).
* Paragraphs are separated by space, never by an indent as well. Butterick again: one or the
  other, not both.
* One family throughout, Archivo, embedded in the file so the document looks the same on a
  machine that has never heard of it. It is the same face the dashboard loads.
* Body ink is #342F49 on white (about 11:1) and the lightest text used at body size is #6C6581
  (about 6:1), both clear of the WCAG 2.2 SC 1.4.3 minimum of 4.5:1. The 3:1 allowance for large
  text is used only for display figures at 18 pt and above.
"""
from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt, RGBColor

import theme as T

FONT = "Archivo"


# ── low-level XML helpers ────────────────────────────────────────────────────
def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(k), str(v))
    return e


def hexc(css):
    return css.lstrip("#").upper()


def rgb(css):
    return RGBColor.from_string(hexc(css))


def _elem(el):
    """Accept a Paragraph, a _Cell or a raw element interchangeably."""
    for attr in ("_p", "_tc", "_element"):
        if hasattr(el, attr):
            return getattr(el, attr)
    return el


def shade(el, css):
    """Solid fill on a paragraph or table cell."""
    el = _elem(el)
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    for old in pr.findall(qn("w:shd")):
        pr.remove(old)
    pr.append(_el("w:shd", **{"w:val": "clear", "w:color": "auto", "w:fill": hexc(css)}))


def borders(el, **sides):
    """sides: top/bottom/left/right -> (size_eighths_pt, colour) or None."""
    el = _elem(el)
    is_cell = el.tag.endswith("}tc")
    pr = el.get_or_add_tcPr() if is_cell else el.get_or_add_pPr()
    tag = "w:tcBorders" if is_cell else "w:pBdr"
    for old in pr.findall(qn(tag)):
        pr.remove(old)
    b = _el(tag)
    for side in ("top", "left", "bottom", "right"):
        spec = sides.get(side)
        if spec is None:
            continue
        size, colour = spec
        b.append(_el(f"w:{side}", **{"w:val": "single", "w:sz": size, "w:space": "0",
                                     "w:color": hexc(colour)}))
    if len(b):
        pr.append(b)


def cell_margins(cell, top=0, left=0, bottom=0, right=0):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcMar")):
        tcPr.remove(old)
    m = _el("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        m.append(_el(f"w:{side}", **{"w:w": int(val * 20), "w:type": "dxa"}))
    tcPr.append(m)


def cell_valign(cell, val="center"):
    cell._tc.get_or_add_tcPr().append(_el("w:vAlign", **{"w:val": val}))


def keep_with_next(p, on=True):
    pPr = p._p.get_or_add_pPr()
    for tag in ("w:keepNext", "w:keepLines"):
        for old in pPr.findall(qn(tag)):
            pPr.remove(old)
    if on:
        pPr.append(_el("w:keepNext"))
        pPr.append(_el("w:keepLines"))


def no_split_row(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(_el("w:cantSplit"))


def repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(_el("w:tblHeader"))


def tracking(run, twentieths):
    """Letter-spacing, in twentieths of a point. Negative tightens display sizes."""
    run._r.get_or_add_rPr().append(_el("w:spacing", **{"w:val": int(twentieths)}))


def tabular_figures(run):
    """Lining, fixed-width numerals so columns of figures line up."""
    rPr = run._r.get_or_add_rPr()
    e = OxmlElement("w14:numSpacing")
    e.set(qn("w14:val"), "tabular")
    rPr.append(e)


# ── the one text primitive everything else is built from ─────────────────────
def para(container, text="", size=10.5, colour=T.N700, bold=False, italic=False,
         space_before=0, space_after=6, leading=1.30, align=None, track=0, caps=False,
         keep=False, style=None):
    p = container.add_paragraph(style=style)
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = leading
    pf.widow_control = True
    if align is not None:
        p.alignment = align
    if text:
        r = add_run(p, text, size=size, colour=colour, bold=bold, italic=italic, track=track,
                    caps=caps)
        del r
    if keep:
        keep_with_next(p)
    return p


def add_run(p, text, size=10.5, colour=T.N700, bold=False, italic=False, track=0, caps=False,
            underline=False):
    # A straight apostrophe in a typeset document is a typewriter artefact, not a character.
    r = p.add_run(text.replace("'", "’"))
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.color.rgb = rgb(colour)
    r.bold = bold
    r.italic = italic
    r.underline = underline
    rPr = r._r.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = _el("w:rFonts")
        rPr.insert(0, rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    if caps:
        rPr.append(_el("w:caps"))
    if track:
        tracking(r, track)
    return r


def rich(container, chunks, **kw):
    """A paragraph built from (text, overrides) pairs.

    An override of {"link": url} renders that run as a real hyperlink, which is how a citation
    sits inside a sentence rather than being exiled to a footnote.
    """
    p = para(container, "", **kw)
    base = dict(size=kw.get("size", 10.5), colour=kw.get("colour", T.N700))
    for text, over in chunks:
        o = dict(base)
        o.update(over or {})
        target = o.pop("link", None)
        if target:
            add_link(p, text, target, size=o.get("size", 10.5),
                     colour=o.get("colour") if o.get("colour") != base["colour"] else None,
                     bold=o.get("bold", False), italic=o.get("italic", False))
        else:
            add_run(p, text, **o)
    return p


# ── page furniture ───────────────────────────────────────────────────────────
PAGE_W, PAGE_H = Cm(21.0), Cm(29.7)
TEXT_W_CM = 15.0


def set_document_defaults(doc):
    """Pin the default run font at the docDefaults level, below every style."""
    styles = doc.styles.element
    dd = styles.find(qn("w:docDefaults"))
    rpd = dd.find(qn("w:rPrDefault"))
    rpr = rpd.find(qn("w:rPr"))
    if rpr is None:
        rpr = _el("w:rPr")
        rpd.append(rpr)
    for old in rpr.findall(qn("w:rFonts")):
        rpr.remove(old)
    rf = _el("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    rpr.insert(0, rf)


def setup_body_section(sec, margin_lr=3.0, margin_top=2.35, margin_bottom=2.15):
    sec.page_width, sec.page_height = PAGE_W, PAGE_H
    sec.left_margin = Cm(margin_lr)
    sec.right_margin = Cm(margin_lr)
    sec.top_margin = Cm(margin_top)
    sec.bottom_margin = Cm(margin_bottom)
    sec.header_distance = Cm(1.25)
    sec.footer_distance = Cm(1.15)
    return sec


def setup_full_bleed_section(sec):
    sec.page_width, sec.page_height = PAGE_W, PAGE_H
    for a in ("left_margin", "right_margin", "top_margin", "bottom_margin",
              "header_distance", "footer_distance"):
        setattr(sec, a, Cm(0))
    return sec


def page_field(p, instr):
    for kind, txt in (("begin", None), ("instrText", instr), ("separate", None),
                      ("t", "1"), ("end", None)):
        if kind == "instrText":
            e = _el("w:instrText", **{"xml:space": "preserve"})
            e.text = txt
        elif kind == "t":
            e = OxmlElement("w:t")
            e.text = txt
        else:
            e = _el("w:fldChar", **{"w:fldCharType": kind})
        r = p.add_run()
        r.font.name = FONT
        r.font.size = Pt(8)
        r.font.color.rgb = rgb(T.N500)
        r._r.append(e)


def build_footer(sec, left_text):
    f = sec.footer
    f.is_linked_to_previous = False
    for p in list(f.paragraphs):
        p._p.getparent().remove(p._p)
    t = f.add_table(rows=1, cols=2, width=Cm(TEXT_W_CM))
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
    for cell, w in zip(t.rows[0].cells, (11.5, 3.5)):
        cell.width = Cm(w)
        cell._tc.get_or_add_tcPr().append(_el("w:tcW", **{"w:w": int(w * 567), "w:type": "dxa"}))
        cell_margins(cell, top=5, bottom=0, left=0, right=0)
        borders(cell, top=(4, T.HAIRLINE))
        pp = cell.paragraphs[0]
        pp.paragraph_format.space_before = Pt(0)
        pp.paragraph_format.space_after = Pt(0)
        pp.paragraph_format.line_spacing = 1.0
    left, right = t.rows[0].cells
    add_run(left.paragraphs[0], left_text, size=7.6, colour=T.N400, track=8, caps=True)
    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_field(rp, " PAGE ")
    tail = f.add_paragraph()
    tail.paragraph_format.space_before = Pt(0)
    tail.paragraph_format.space_after = Pt(0)
    tail.paragraph_format.line_spacing = 1.0
    add_run(tail, "", size=1)


def build_header(sec, text):
    h = sec.header
    h.is_linked_to_previous = False
    p = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(9)
    add_run(p, text, size=7.6, colour=T.N400, track=8, caps=True)


def full_page_image(container, path):
    """A page-anchored, behind-text image at exactly A4, so it cannot push onto a second page.

    Built by rewriting the wp:inline python-docx produces into a wp:anchor. Element order in
    wp:anchor is fixed by the schema, so the XML is assembled as text rather than by appending.
    """
    from docx.oxml import parse_xml

    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(path, width=PAGE_W, height=PAGE_H)
    drawing = run._r.find(qn("w:drawing"))
    inline = drawing[0]

    NSWP = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
    NSA = "http://schemas.openxmlformats.org/drawingml/2006/main"
    graphic = inline.find("{%s}graphic" % NSA)
    docPr = inline.find("{%s}docPr" % NSWP)
    doc_id = docPr.get("id")
    doc_name = docPr.get("name")
    from lxml import etree
    graphic_xml = etree.tostring(graphic, encoding="unicode")

    anchor_xml = (
        '<wp:anchor xmlns:wp="%s" xmlns:a="%s" distT="0" distB="0" distL="0" distR="0" '
        'simplePos="0" relativeHeight="1" behindDoc="1" locked="0" layoutInCell="1" '
        'allowOverlap="1">'
        '<wp:simplePos x="0" y="0"/>'
        '<wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>'
        '<wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>'
        '<wp:extent cx="%d" cy="%d"/>'
        '<wp:effectExtent l="0" t="0" r="0" b="0"/>'
        '<wp:wrapNone/>'
        '<wp:docPr id="%s" name="%s"/>'
        '<wp:cNvGraphicFramePr/>'
        '%s'
        '</wp:anchor>' % (NSWP, NSA, int(PAGE_W), int(PAGE_H), doc_id, doc_name, graphic_xml)
    )
    drawing.remove(inline)
    drawing.append(parse_xml(anchor_xml))
    return p


# ── headings ─────────────────────────────────────────────────────────────────
def section_heading(doc, number, title, standfirst=None, first=False):
    """Sections flow rather than each claiming a fresh page.

    A forced page break before every section leaves a stub page whenever a section overruns by a
    line or two, and a page holding one paragraph reads as a mistake. Instead the opener carries a
    rule above it so a section start is unmistakable mid-page, and the kicker, title, standfirst
    and first paragraph are bound together so the break can never fall inside them.
    """
    if not first:
        rule = para(doc, "", size=1, space_before=16, space_after=0, keep=True)
        borders(rule, top=(10, T.BRAND_LINE))
    p = para(doc, "", space_before=0 if first else 12, space_after=2, keep=True)
    add_run(p, number, size=8.5, colour=T.BRAND, bold=True, track=26, caps=True)
    h = para(doc, "", space_before=0, space_after=3 if standfirst else 8, leading=1.10, keep=True)
    r = add_run(h, title, size=19, colour=T.BRAND_DEEP, bold=True, track=-8)
    del r
    borders(h, bottom=None)
    if standfirst:
        s = para(doc, standfirst, size=11.5, colour=T.N500, space_before=0, space_after=9,
                 leading=1.28, keep=True)
        del s
    return h


def h2(doc, text, num=None):
    p = para(doc, "", space_before=12, space_after=3, leading=1.20, keep=True)
    if num:
        add_run(p, num + "  ", size=13, colour=T.BRAND, bold=True, track=-5)
    add_run(p, text, size=13, colour=T.N800, bold=True, track=-5)
    return p


def h3(doc, text):
    p = para(doc, "", space_before=9, space_after=2, leading=1.22, keep=True)
    add_run(p, text, size=10.5, colour=T.N800, bold=True, track=-2)
    return p


def label(doc, text, colour=None):
    p = para(doc, "", space_before=10, space_after=3, keep=True)
    add_run(p, text, size=8.0, colour=colour or T.BRAND, bold=True, track=24, caps=True)
    return p


def bullets(doc, items, size=10.5, space_after=4, colour=None):
    for it in items:
        p = para(doc, "", size=size, space_after=space_after, leading=1.30)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.55)
        add_run(p, "•   ", size=size, colour=T.BRAND_BRIGHT, bold=True)
        if isinstance(it, tuple):
            add_run(p, it[0], size=size, colour=colour or T.N800, bold=True)
            add_run(p, it[1], size=size, colour=colour or T.N700)
        else:
            add_run(p, it, size=size, colour=colour or T.N700)


# ── tables ───────────────────────────────────────────────────────────────────
# House style: no vertical rules at all, one solid header band, hairline row
# separators, numerals right-aligned and tabular. Rules that separate rather
# than enclose. A grid of boxes reads slower than a set of aligned rows.
def table(doc, headers, rows, widths, aligns=None, zebra=False, header_fill=None,
          size=9.3, head_size=7.8, pad=5.0, first_bold=False, note=None):
    aligns = aligns or ["l"] * len(headers)
    t = doc.add_table(rows=0, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))

    total = sum(widths)
    cm = [TEXT_W_CM * w / total for w in widths]

    def put(cells, values, is_head):
        for i, (cell, val) in enumerate(zip(cells, values)):
            cell.width = Cm(cm[i])
            cell._tc.get_or_add_tcPr().append(
                _el("w:tcW", **{"w:w": int(cm[i] * 567), "w:type": "dxa"}))
            p = cell.paragraphs[0]
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = 1.15 if is_head else 1.22
            a = aligns[i]
            if a == "r":
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif a == "c":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            chunks = val if isinstance(val, list) else [(str(val), None)]
            for text, over in chunks:
                o = dict(size=head_size if is_head else size,
                         colour="#FFFFFF" if is_head else T.N700,
                         bold=is_head or (first_bold and i == 0 and not is_head))
                if is_head:
                    o.update(track=14, caps=True)
                o.update(over or {})
                target = o.pop("link", None)
                if target:
                    add_link(p, text, target, size=o.get("size", size),
                             bold=o.get("bold", False), italic=o.get("italic", False))
                    continue
                r = add_run(p, text, **o)
                if a in ("r", "c"):
                    tabular_figures(r)
            cell_margins(cell, top=pad, bottom=pad, left=6.5 if i else 0, right=6.5)
            cell_valign(cell, "center" if is_head else "top")

    hrow = t.add_row()
    put(hrow.cells, headers, True)
    for c in hrow.cells:
        shade(c._tc, header_fill or T.BRAND_DEEP)
        borders(c._tc)
    repeat_header(hrow)
    no_split_row(hrow)
    for c in hrow.cells:
        for pp in c.paragraphs:
            keep_with_next(pp)

    for n, values in enumerate(rows):
        row = t.add_row()
        put(row.cells, values, False)
        last = n == len(rows) - 1
        for c in row.cells:
            if zebra and n % 2 == 1:
                shade(c._tc, T.N50)
            borders(c._tc, bottom=(6, T.N300 if last else T.N200))
        no_split_row(row)

    if note:
        para(doc, note, size=8.2, colour=T.N500, space_before=5, space_after=9, leading=1.28)
    else:
        para(doc, "", size=1, space_after=4)
    return t


def figure(doc, path, caption, width_cm=TEXT_W_CM, space_before=8, space_after=0):
    p = para(doc, "", space_before=space_before, space_after=4, leading=1.0, keep=True)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(path, width=Cm(width_cm))
    c = para(doc, "", size=8.2, colour=T.N500, space_before=5,
             space_after=space_after + 8, leading=1.30)
    add_run(c, caption, size=8.2, colour=T.N500)
    borders(c, top=(4, T.N200))
    return p


# ── panels ───────────────────────────────────────────────────────────────────
def panel(doc, title, body, accent=None, fill=None, title_colour=None, space_after=9):
    """A single-cell table used as a bordered note. The left rule carries the emphasis."""
    accent = accent or T.BRAND
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
    cell = t.rows[0].cells[0]
    cell.width = Cm(TEXT_W_CM)
    cell._tc.get_or_add_tcPr().append(_el("w:tcW", **{"w:w": int(TEXT_W_CM * 567),
                                                      "w:type": "dxa"}))
    shade(cell._tc, fill or T.BRAND_SOFT)
    borders(cell._tc, left=(24, accent))
    cell_margins(cell, top=8, bottom=8, left=10, right=10)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4) if body else Pt(0)
    p.paragraph_format.line_spacing = 1.2
    if title:
        add_run(p, title, size=8.0, colour=title_colour or accent, bold=True, track=22, caps=True)
    if body:
        bp = cell.add_paragraph()
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(0)
        bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        bp.paragraph_format.line_spacing = 1.32
        chunks = body if isinstance(body, list) else [(body, None)]
        for text, over in chunks:
            o = dict(size=9.8, colour=T.N700)
            o.update(over or {})
            target = o.pop("link", None)
            if target:
                add_link(bp, text, target, size=o.get("size", 9.8), bold=o.get("bold", False),
                         italic=o.get("italic", False))
            else:
                add_run(bp, text, **o)
    no_split_row(t.rows[0])
    para(doc, "", size=1, space_before=0, space_after=space_after)
    return t


def stat_band(doc, stats, space_after=9):
    """Display figures across the measure, each with a caption and a source line."""
    t = doc.add_table(rows=1, cols=len(stats))
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
    w = TEXT_W_CM / len(stats)
    for i, (value, cap, src) in enumerate(stats):
        cell = t.rows[0].cells[i]
        cell.width = Cm(w)
        cell._tc.get_or_add_tcPr().append(_el("w:tcW", **{"w:w": int(w * 567), "w:type": "dxa"}))
        borders(cell._tc, top=(18, T.BRAND))
        cell_margins(cell, top=7, bottom=6, left=0 if i == 0 else 5, right=7)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.0
        r = add_run(p, value, size=19, colour=T.BRAND_DEEP, bold=True, track=-16)
        tabular_figures(r)
        c = cell.add_paragraph()
        c.paragraph_format.space_before = Pt(0)
        c.paragraph_format.space_after = Pt(2)
        c.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        c.paragraph_format.line_spacing = 1.24
        add_run(c, cap, size=8.2, colour=T.N700)
        s = cell.add_paragraph()
        s.paragraph_format.space_before = Pt(0)
        s.paragraph_format.space_after = Pt(0)
        s.paragraph_format.line_spacing = 1.2
        if isinstance(src, tuple):
            add_link(s, src[0], src[1], size=7.2, colour=T.BRAND)
        else:
            add_run(s, src, size=7.2, colour=T.N400)
    no_split_row(t.rows[0])
    para(doc, "", size=1, space_after=space_after)
    return t


def code_block(doc, lines, caption=None):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", **{"w:type": "fixed"}))
    cell = t.rows[0].cells[0]
    cell.width = Cm(TEXT_W_CM)
    cell._tc.get_or_add_tcPr().append(_el("w:tcW", **{"w:w": int(TEXT_W_CM * 567),
                                                      "w:type": "dxa"}))
    shade(cell._tc, T.RAIL)
    cell_margins(cell, top=9, bottom=9, left=10, right=10)
    first = True
    for text, colour in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        p.paragraph_format.line_spacing = 1.30
        r = p.add_run(text)
        r.font.size = Pt(8.2)
        r.font.color.rgb = rgb(colour)
        rPr = r._r.get_or_add_rPr()
        rf = _el("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rf.set(qn(a), "Consolas")
        rPr.insert(0, rf)
    no_split_row(t.rows[0])
    if caption:
        c = para(doc, caption, size=8.2, colour=T.N500, space_before=5, space_after=9,
                 leading=1.28)
        borders(c, top=(4, T.N200))
    else:
        para(doc, "", size=1, space_after=9)
    return t


def page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = add_run(p, "", size=1)
    r.add_break(WD_BREAK.PAGE)
    return p


# ── hyperlinks and citations ─────────────────────────────────────────────────
def add_link(p, text, target, size=10.5, colour=None, bold=False, italic=False):
    """A real external hyperlink. python-docx has no API for these, so the relationship is
    registered on the part and a w:hyperlink element is written by hand."""
    part = p.part
    r_id = part.relate_to(
        target,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = _el("w:hyperlink", **{"r:id": r_id})
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = _el("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), FONT)
    rPr.append(rf)
    rPr.append(_el("w:color", **{"w:val": hexc(colour or T.BRAND)}))
    rPr.append(_el("w:sz", **{"w:val": int(size * 2)}))
    rPr.append(_el("w:szCs", **{"w:val": int(size * 2)}))
    if bold:
        rPr.append(OxmlElement("w:b"))
    if italic:
        rPr.append(OxmlElement("w:i"))
    rPr.append(_el("w:u", **{"w:val": "single", "w:color": hexc(colour or T.BRAND_LINE)}))
    run.append(rPr)
    t = _el("w:t", **{"xml:space": "preserve"})
    t.text = text.replace("'", "\u2019")
    run.append(t)
    link.append(run)
    p._p.append(link)
    return link


def source_line(container, chunks, size=8.0, space_before=3, space_after=9):
    """The small grey attribution that sits under a figure or a set of figures.

    `chunks` is a list of (label, url) pairs; a None url renders as plain text.
    """
    p = para(container, "", size=size, colour=T.N500, space_before=space_before,
             space_after=space_after, leading=1.28)
    add_run(p, "Source: ", size=size, colour=T.N400)
    for i, (label, target) in enumerate(chunks):
        if i:
            add_run(p, "  ·  ", size=size, colour=T.N300)
        if target:
            add_link(p, label, target, size=size, colour=T.BRAND)
        else:
            add_run(p, label, size=size, colour=T.N500)
    return p
